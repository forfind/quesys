# 导入模块
from io import BytesIO

from PIL import Image
from docx import Document
# 此模块中包含 docx 中各类单位方法
from docx import shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def change_to_img(t,q,img_data,n):
    print(type(img_data))
    # 将字节对象转为Byte字节流数据,供Image.open使用
    byte_stream = BytesIO(img_data)
    print(type(byte_stream))
    roiImg = Image.open(byte_stream)
    # 图片保存
    img_path = "./output/"+str(t)+ "_" + str(q)+ "_"  + str(n) + '.png'
    roiImg.save(img_path)
    #with open(img_path, 'wb') as f:
    #    f.write(imgByteArr)
    return img_path

def write_TestandAns(code,list):
    doc = Document()
    doc.add_heading('【试题'+str(code)+"】（含答案）",1)
    cate = ""
    i = 0
    numlist = ["一", "二", "三", "四", "五", "六"]
    for ques in list:
        if cate != ques[0]:
            cate = ques[0]
            doc.add_heading(numlist[i]+"、"+cate,2)
            i += 1
        doc.add_paragraph(ques[2]+"."+"["+ques[1]+"分]"+ques[3])
        if str(ques[4])!="None":
            picpath = change_to_img(code,int(ques[2]),ques[4],1)
            paragraph = doc.add_paragraph()
            #图片居中设置
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("")
            run.add_picture(picpath)
        doc.add_paragraph("答案："+ques[5])
        if str(ques[6])!="None":
            picpath = change_to_img(code,int(ques[2]),ques[6],2)
            paragraph = doc.add_paragraph()
            #图片居中设置
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("")
            run.add_picture(picpath)

    # 保存文件
    doc.save('./output/试题'+str(code)+'(含答案)'+'.docx')

def write_Test(code,list):
    doc = Document()
    doc.add_heading('【试题'+str(code)+"】（仅试题）",1)
    cate = ""
    i = 0
    numlist = ["一", "二", "三", "四", "五", "六"]
    for ques in list:
        if cate != ques[0]:
            cate = ques[0]
            doc.add_heading(numlist[i]+"、"+cate,2)
            i += 1
        doc.add_paragraph(ques[2]+"."+"["+ques[1]+"分]"+ques[3])
        if str(ques[4])!="None":
            picpath = change_to_img(code,int(ques[2]),ques[4],1)
            paragraph = doc.add_paragraph()
            #图片居中设置
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("")
            run.add_picture(picpath)
    # 保存文件
    doc.save('./output/试题'+str(code)+'(仅试题)'+'.docx')

def main():
    fin = open("8.png", 'rb')
    print(fin.read())
    img = fin.read()
    print(type(img))
    fin.close()
    change_to_img(1,1,img,0)

if __name__ == "__main__":
    main()